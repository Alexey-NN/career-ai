from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_line(doc):
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2b, 0x57, 0x9a)
    run.font.name = 'Calibri'

def add_job_heading(doc, title, company, dates, location, italic_note=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.space_after = Pt(1)
    run = p.add_run(company + " · " + dates + " · " + location)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = 'Calibri'

    if italic_note:
        p = doc.add_paragraph()
        p.space_after = Pt(3)
        run = p.add_run(italic_note)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        run.italic = True
        run.font.name = 'Calibri'

def add_scope(doc, text):
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    run = p.add_run("Зона ответственности: ")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'

def add_results_label(doc):
    p = doc.add_paragraph()
    p.space_before = Pt(3)
    p.space_after = Pt(2)
    run = p.add_run("Результаты:")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'

def add_bullet(doc, bold_part, rest):
    p = doc.add_paragraph(style='List Bullet')
    p.space_after = Pt(2)
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
    if rest:
        run = p.add_run(rest)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'

def add_sub_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet 2')
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

# ===================== HEADER =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Алексей Козлов")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
run.font.name = 'Calibri'
p.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("COO | Операционный директор")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x2b, 0x57, 0x9a)
run.font.name = 'Calibri'
p.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("Операционная трансформация · P&L · Масштабирование · Стартапы и международные холдинги")
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.italic = True
p.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("Москва, Россия · Не готов к переезду, готов к командировкам")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.space_after = Pt(1)

p = doc.add_paragraph()
run = p.add_run("Telegram: @Alexey_Kozlov_NN · Email: ak_nn@mail.ru · Телефон: +7 (903) 1248453")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.space_after = Pt(10)

# ===================== ABOUT =====================
add_line(doc)
add_section_heading(doc, "Обо мне")

p = doc.add_paragraph()
p.space_after = Pt(6)
run = p.add_run("Операционный руководитель с 20+ годами управления P&L, кросс-функциональными командами и трансформациями в международных холдингах. Три раза строил с нуля — IT-консалтинг, Shared Service Center, направление внедрения ERP. Умею переводить стратегию в процессы, которые масштабируются без потери контроля. В кризисе не замораживаю бизнес — перестраиваю: −15% рынка обернул в +15% маржинальности.")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'

# ===================== EXPERIENCE =====================
add_line(doc)
add_section_heading(doc, "Опыт работы")

# CEO
add_job_heading(doc,
    "Генеральный директор, член совета директоров (CEO)",
    "SCHNEIDER GROUP",
    "Июль 2022 — Февраль 2026",
    "Москва",
    "Топ-10 компаний РФ в аутсорсинге учётных функций. 150+ сотрудников, 8 бухгалтерских групп, IT-отдел, 3 региональных хаба (Москва, Берлин, Алматы)."
)

add_scope(doc, "Полное управление бизнес-направлениями аутсорсинга и IT — стратегия, P&L, бюджетирование, развитие команды. Взаимодействие с Советом директоров и собственником, защита бюджетов.")

add_results_label(doc)
add_bullet(doc, "Антикризис: ", "В условиях −15% выручки и санкций — оптимизация штата −32% с удержанием ключевых экспертов, стабилизация клиентской базы")
add_bullet(doc, "Финансовая эффективность: ", "−30% прямых затрат, +15% Direct margin, +45% маржинальность на сотрудника — через процессный подход и P&L-дисциплину")
add_bullet(doc, "Масштабируемость: ", "Унифицировал стандарты в 8 группах и 3 хабах, внедрил процессный подход — рост без потери управляемости")
add_bullet(doc, "Коммерция: ", "Привлечение клиентов через тендеры и КП, долгосрочные отношения с ключевыми заказчиками")
add_bullet(doc, "Цифровизация: ", "Развитие направления внедрения 1С-решений, повышение операционной эффективности клиентов")

# Director Shared Services
add_job_heading(doc,
    "Исполнительный директор (Director Shared Services)",
    "WORTMANN GROUP",
    "Сентябрь 2015 — Июнь 2022",
    "Москва",
    "Немецкий холдинг — производство и дистрибуция обуви. 5 юрлиц в РФ. Прямое подчинение CFO холдинга (Германия)."
)

add_scope(doc, "Общие сервисы для группы — бухгалтерия, ДЗ, IT, логистика, юристы, админ. Разработка операционной стратегии, управленческая отчётность, защита бюджетов. (CEO/CFO/COO-уровень)")

add_results_label(doc)
add_bullet(doc, "Оборотный капитал: ", "Создал отдел ДЗ с нуля — кредитные лимиты, блокировки/разблокировки отгрузок, мониторинг. DSO −25%, высвобождение оборотных средств")
add_bullet(doc, "«Честный ЗНАК»: ", "Возглавил внедрение маркировки обуви (2020–2021) — интеграция с 3PL, ЭДО, клиентскими системами. Пилот Минпромторга без штрафов и простоев")
add_bullet(doc, "Логистика: ", "Пересмотрел тарифную модель 3PL (оплата за короб), консолидировал отгрузки, внедрил имплантов ТК на склад — −30% логистических издержек при росте предсказуемости")
add_bullet(doc, "Shared Service Center: ", "Объединил бэк-офисы 5 юрлиц в единый SSC с прозрачными процессами и прямым подчинением CFO холдинга")
add_bullet(doc, "Комплаенс: ", "Выполнение требований немецкого головного офиса и российского законодательства в условиях меняющегося регулирования")

# CIO
add_job_heading(doc,
    "Директор по информационным технологиям (CIO)",
    "SCHNEIDER GROUP",
    "Ноябрь 2005 — Август 2015",
    "Москва",
    "30 специалистов, 5 групп: инфраструктура, техподдержка, ИБ, консалтинг. 2 ЦОД, 6 удалённых офисов, 400+ баз данных, 450+ пользователей."
)

add_scope(doc, "IT-инфраструктура, информационная безопасность, коммерческое направление (1С, MS Dynamics NAV, SAP), технологическая стратегия.")

add_results_label(doc)
add_bullet(doc, "Инфраструктура: ", "Отказоустойчивая архитектура (2 ЦОД + 6 офисов, 400+ БД) — минимизация простоев и рисков")
add_bullet(doc, "ИБ: ", "Комплекс мер по защите данных — соответствие требованиям международных клиентов и РФ, снижение рисков утечек и штрафов")
add_bullet(doc, "Новый бизнес: ", "Создал с нуля направление IT-консалтинга и внедрения — от пресейла до сдачи проектов. Новая бизнес-линия компании")
add_bullet(doc, "Ключевые проекты:", "")
add_sub_bullet(doc, "1С для международных производителей одежды и обуви — интеграция с системами головных компаний")
add_sub_bullet(doc, "1С + ИТАН + SAP для мирового производителя игрушек — унификация отчётности для головного офиса")
add_sub_bullet(doc, "1С + логистические системы для японского автопроизводителя — прозрачность цепочки поставок")
add_bullet(doc, "EDI и МСФО: ", "Внедрил EDI и системы международного учёта (IFRS/US GAAP) — ускорение консолидации и принятия решений головными офисами")

# ===================== COMPETENCIES =====================
add_line(doc)
add_section_heading(doc, "Ключевые компетенции для COO")

table = doc.add_table(rows=7, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light List Accent 1'

data = [
    ("Операционное управление", "P&L, процессный подход, масштабирование, контроль исполнения"),
    ("Трансформация", "Антикризисное управление, реструктуризация, построение SSC с нуля"),
    ("Финансы", "Бюджетирование, управление оборотным капиталом, DSO, cost optimization"),
    ("Логистика и цепочки поставок", "3PL, маркировка, консолидация отгрузок"),
    ("IT и цифровизация", "1С, SAP, MS Dynamics NAV, ИБ, EDI, МСФО/US GAAP"),
    ("Команда", "150+ сотрудников, кросс-функциональные команды, удержание в кризис"),
    ("Комплаенс", "Российское законодательство, требования международных HQ"),
]

for i, (area, skills) in enumerate(data):
    row = table.rows[i]
    cell0 = row.cells[0]
    cell1 = row.cells[1]
    p0 = cell0.paragraphs[0]
    run = p0.add_run(area)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p1 = cell1.paragraphs[0]
    run = p1.add_run(skills)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

# ===================== EDUCATION =====================
add_line(doc)
add_section_heading(doc, "Образование")

ed_items = [
    ("CEO", "Eduson Academy", "2026"),
    ("Переговорный практикум «ДоговоРинг»", "Школа Переговоров Галины Жуковой", "2024"),
    ("Персональное управленческое искусство", "Таллинская школа менеджеров", "2019"),
    ("Президентская программа управленческих кадров, Финансы и кредит", "Нижегородский государственный технический университет", "2004"),
    ("Стажировка в Италии, Tacis MTP", "TACIS (EC), IFOA (Италия)", "2004"),
    ("Экономика и управление на предприятии", "Волжская государственная академия водного транспорта", "2000"),
    ("Электронные вычислительные машины", "Нижегородский политехнический институт", "1992"),
]

for title, school, year in ed_items:
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run = p.add_run(" — " + school + " (" + year + ")")
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'

# ===================== LANGUAGES =====================
add_line(doc)
add_section_heading(doc, "Языки")

p = doc.add_paragraph()
p.space_after = Pt(2)
run = p.add_run("Русский")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
run = p.add_run(" — родной")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.space_after = Pt(2)
run = p.add_run("Английский")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
run = p.add_run(" — C1, продвинутый (переговоры и документация)")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'

# Save
path = r"C:\Users\Alexey Kozlov\career-ai\data\vacancies\COO_Startup\260507_CV_COO_Startup_RU.docx"
doc.save(path)
print("Saved:", path)